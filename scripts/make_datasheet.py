"""
Generate the MnemoScript product datasheet as a styled .docx.

Run:  python scripts/make_datasheet.py
Output: MnemoScript_Datasheet.docx in the repo root.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── palette ────────────────────────────────────────────────────────────────
INK        = RGBColor(0x1E, 0x29, 0x3B)   # slate-800 body text
ACCENT     = RGBColor(0x4F, 0x46, 0xE5)   # indigo / periwinkle
ACCENT_HEX = "4F46E5"
DEEP_HEX   = "312E81"                       # indigo-900 banners
MUTED      = RGBColor(0x64, 0x74, 0x8B)   # slate-500
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_HEX  = "EEF0FB"                       # light periwinkle row fill
RULE_HEX   = "C7D2FE"                       # soft border
BODY_FONT  = "Segoe UI"

# ── low-level helpers ──────────────────────────────────────────────────────
def _set(el, **kw):
    for k, v in kw.items():
        el.set(qn(k), v)

def shade(cell, hex_fill):
    shd = OxmlElement('w:shd')
    _set(shd, **{'w:val': 'clear', 'w:color': 'auto', 'w:fill': hex_fill})
    cell._tc.get_or_add_tcPr().append(shd)

def cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('bottom', bottom), ('start', left), ('end', right)):
        e = OxmlElement(f'w:{tag}')
        _set(e, **{'w:w': str(val), 'w:type': 'dxa'})
        m.append(e)
    tcPr.append(m)

def vcenter(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement('w:vAlign')
    _set(va, **{'w:val': 'center'})
    tcPr.append(va)

def table_borders(table, color=RULE_HEX, inside=True, box=False):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    edges = ['insideH'] if inside else []
    if box:
        edges += ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']
    for edge in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
        e = OxmlElement(f'w:{edge}')
        if edge in edges:
            _set(e, **{'w:val': 'single', 'w:sz': '4', 'w:space': '0', 'w:color': color})
        else:
            _set(e, **{'w:val': 'none', 'w:sz': '0', 'w:space': '0', 'w:color': 'auto'})
        borders.append(e)
    tblPr.append(borders)

def set_col_widths(table, widths):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

def para_border_bottom(p, color=ACCENT_HEX, sz=12, space=6):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    _set(b, **{'w:val': 'single', 'w:sz': str(sz), 'w:space': str(space), 'w:color': color})
    pbdr.append(b)
    pPr.append(pbdr)

def run(p, text, *, size=10, bold=False, italic=False, color=INK, font=BODY_FONT, caps=False):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    if caps:
        r.font.all_caps = True
        rPr = r._element.get_or_add_rPr()
        sp = OxmlElement('w:spacing'); _set(sp, **{'w:val': '30'}); rPr.append(sp)
    return r

def space_after(p, pts):
    p.paragraph_format.space_after = Pt(pts)
    return p

# ── document build ─────────────────────────────────────────────────────────
doc = Document()
normal = doc.styles['Normal']
normal.font.name = BODY_FONT
normal.font.size = Pt(10)
normal.font.color.rgb = INK

sec = doc.sections[0]
sec.top_margin = Inches(0.7)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.8)
sec.right_margin = Inches(0.8)
CONTENT_W = Inches(6.9)

def section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run(p, text, size=13.5, bold=True, color=ACCENT, caps=True)
    para_border_bottom(p)
    return p

def lead(text):
    p = doc.add_paragraph()
    space_after(p, 8)
    run(p, text, size=10.5, color=INK)
    return p

def feature_table(items):
    """items: list of (name, description)."""
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_borders(t, inside=True)
    set_col_widths(t, [Inches(2.05), Inches(4.85)])
    # header
    h = t.rows[0].cells
    for c in h:
        shade(c, DEEP_HEX); cell_margins(c); vcenter(c)
    h[0].paragraphs[0].clear() if False else None
    run(h[0].paragraphs[0], "Feature", size=9, bold=True, color=WHITE, caps=True)
    run(h[1].paragraphs[0], "What it does", size=9, bold=True, color=WHITE, caps=True)
    # rows
    for i, (name, desc) in enumerate(items):
        row = t.add_row().cells
        for c in row:
            cell_margins(c)
        if i % 2 == 1:
            shade(row[0], LIGHT_HEX); shade(row[1], LIGHT_HEX)
        run(row[0].paragraphs[0], name, size=9.5, bold=True, color=ACCENT)
        run(row[1].paragraphs[0], desc, size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ── 1. Title banner ────────────────────────────────────────────────────────
banner = doc.add_table(rows=1, cols=1)
banner.alignment = WD_TABLE_ALIGNMENT.CENTER
set_col_widths(banner, [CONTENT_W])
bc = banner.rows[0].cells[0]
shade(bc, DEEP_HEX)
cell_margins(bc, top=220, bottom=220, left=260, right=260)
p = bc.paragraphs[0]
space_after(p, 2)
run(p, "MnemoScript", size=30, bold=True, color=WHITE)
p2 = bc.add_paragraph(); space_after(p2, 2)
run(p2, "The All-in-One Writing & Worldbuilding Studio", size=12.5, color=RGBColor(0xC7, 0xD2, 0xFE))
p3 = bc.add_paragraph()
run(p3, "PRODUCT DATASHEET", size=8.5, bold=True, color=RGBColor(0xA5, 0xB4, 0xFC), caps=True)
run(p3, "      •      Revision: June 2026      •      Status: Active Preview",
    size=8.5, color=RGBColor(0xA5, 0xB4, 0xFC))
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── 2. Overview ────────────────────────────────────────────────────────────
section_heading("At a Glance")
lead("MnemoScript is a desktop studio that unifies long-form writing, visual brainstorming, "
     "and fantasy cartography in one offline-first workspace. Authors, worldbuilders, game "
     "masters, and screenwriters can draft chapters, map out plots, and illustrate entire "
     "worlds without ever leaving the application — or handing their work to the cloud. "
     "Everything is stored as plain files on your own disk, so your projects stay private, "
     "portable, and yours.")

# Highlights strip (3 cards)
cards = [
    ("3-in-1 workspace", "Rich-text Notes, Mind Maps, and a Fantasy Map studio share one project."),
    ("Offline & private", "All content lives on your device as plain files — no account, no cloud."),
    ("Cross-platform", "Runs on Windows desktop and Android, built on a fast, secure Tauri shell."),
]
ct = doc.add_table(rows=1, cols=3)
ct.alignment = WD_TABLE_ALIGNMENT.CENTER
set_col_widths(ct, [Inches(2.3)] * 3)
for i, (title, body) in enumerate(cards):
    c = ct.rows[0].cells[i]
    shade(c, LIGHT_HEX); cell_margins(c, top=140, bottom=140, left=160, right=160)
    tp = c.paragraphs[0]; space_after(tp, 3)
    run(tp, title, size=10.5, bold=True, color=ACCENT)
    bp = c.add_paragraph()
    run(bp, body, size=9, color=INK)
table_borders(ct, inside=False)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── 3. Document types ──────────────────────────────────────────────────────
section_heading("Three Document Types, One Project")
lead("Create any mix of these inside a single project folder and switch between them instantly.")
feature_table([
    ("Rich-Text Notes",
     "A distraction-free word processor for chapters, scenes, lore, and outlines — with "
     "formatting, images, grammar checking, and to-do lists."),
    ("Mind Maps",
     "A draw.io-style canvas for mapping plots, character webs, and world relationships with "
     "draggable nodes and connectors."),
    ("Fantasy Map Studio",
     "A full cartography engine that generates and hand-paints illustrated fantasy maps in "
     "either a hand-drawn ink or classic colored style."),
])

# ── 4. Writing & editor ────────────────────────────────────────────────────
section_heading("Writing & Rich-Text Editor")
feature_table([
    ("Focused editor", "A clean writing surface (TipTap / ProseMirror) with a smooth animated caret that keeps you in flow."),
    ("Full formatting", "Bold, italic, underline, three heading levels, bullet & numbered lists, block quotes, code blocks, and dividers."),
    ("Text alignment", "Left, center, right, and justified paragraphs for finished-looking pages."),
    ("Slash “/” menu", "A Notion-style command palette — type “/” to insert any block without reaching for the mouse."),
    ("Image embedding", "“/image” opens a native picker; the picture is copied into the project and rendered inline, travelling with your files."),
    ("Grammar & spelling", "Real-time linguistic checking (LanguageTool) underlines issues and offers one-click corrections."),
    ("Live word count", "Always-on word and character counts in the status bar to track your progress."),
])

# ── 5. To-do (the new feature, highlighted) ────────────────────────────────
section_heading("To-Do Lists & Task Tracking")
lead("New: a built-in task manager that lives right inside your notes — perfect for revision "
     "checklists, research leads, and chapter goals. Available in Notes only, so it never "
     "clutters your maps.")
feature_table([
    ("“/todo” command", "Type “/todo” to drop a classic checkbox to-do list straight into a note."),
    ("Check off tasks", "Click a checkbox to mark a task done — completed items strike through and gently fade."),
    ("Nested sub-tasks", "Press Tab to indent a task beneath another, building multi-level checklists."),
    ("Sidebar shortcut", "A dedicated “To-do List” button in the Notes side panel, shown only for notes."),
])

# ── 6. Organization ────────────────────────────────────────────────────────
section_heading("Organization & Project Management")
feature_table([
    ("Projects on disk", "Each project is a real folder with one file per document plus an assets folder — fully portable, no database."),
    ("Nestable folders", "Build a directory tree of any depth to organize chapters, notes, and references."),
    ("Drag-and-drop", "Reorder and move documents or folders by dragging, or via a right-click “Move to” menu."),
    ("Auto-save + manual", "Configurable auto-save (15 / 30 / 60s) and one-click save, with a live synced/unsaved indicator."),
])

# ── 7. Mind mapping ────────────────────────────────────────────────────────
section_heading("Mind Mapping")
feature_table([
    ("Visual canvas", "A draw.io-style board (React Flow) for brainstorming plots, characters, and world connections."),
    ("Nodes & connectors", "Draggable nodes, drawn links, colors, and free pan / zoom navigation."),
])

# ── 8. Fantasy map studio ──────────────────────────────────────────────────
section_heading("Fantasy Map Studio")
lead("A complete map-making engine — generate a world in one click, then refine every coastline, "
     "mountain, and border by hand.")
feature_table([
    ("Two art styles", "Switch any map between a hand-drawn ink / parchment look and a classic colored, Inkarnate-style render — losing nothing."),
    ("Procedural generation", "Auto-create continents, islands, or archipelagos with controls for land amount, coastline detail, biomes, and icon density."),
    ("Paintable terrain", "Raise land or carve water with size-adjustable brushes on a shared heightmap; a live brush ring follows the cursor."),
    ("Water features", "Dedicated Sea, Lake, and River modes paint connected, coastline-aware water with ripples."),
    ("Ruggedness control", "A single slider warps coastlines from smooth to fractal — bays, peninsulas, and fjords — live on any map."),
    ("Hand-drawn icons", "An original ink icon set (mountains, forests, hamlets to great cities, castles, dungeons, temples) plus ~60 curated glyphs."),
    ("Scatter brush", "Drag to paint many icons at once for forests, hills, and clustered settlements."),
    ("Regions & borders", "Draw dotted region borders that close at the start point, then nest realms → provinces → counties with auto-parenting."),
    ("Smart overlap", "Front assets cleanly hide those behind them, so overlapping mountains read as a single range."),
    ("Naming & labels", "Double-click any asset to name it, with red serif place-names in the hand-drawn style."),
    ("Big canvases & import", "Build maps up to 4K, import your own art (PNG / SVG / WebP), and export the result to PNG."),
])

# ── 9. Publishing & personalization ────────────────────────────────────────
section_heading("Publishing & Export")
feature_table([
    ("Compile to PDF book", "Combine chapters into one book — set title & author, reorder chapters, toggle a cover and table of contents, then export to PDF."),
    ("Map export", "Save any fantasy map to a high-resolution PNG image."),
])

section_heading("Personalization")
feature_table([
    ("Six themes", "Midnight, Parchment, Nebula, Ocean, Forest, and Sunset color profiles for any mood or lighting."),
    ("Typography controls", "Tune font family, size, line spacing, and page margins to your taste."),
])

# ── 10. Technical specifications ───────────────────────────────────────────
section_heading("Technical Specifications")
specs = [
    ("Product", "MnemoScript — writing & worldbuilding studio"),
    ("Platforms", "Windows (desktop) · Android"),
    ("Application shell", "Tauri 2 (Rust backend + native system WebView)"),
    ("Frontend", "React 19 · TypeScript · Vite 8 · Tailwind CSS v4"),
    ("Rich-text engine", "TipTap 3 (ProseMirror)"),
    ("Mind-map engine", "React Flow (@xyflow/react 12)"),
    ("Map engine", "Konva / react-konva · simplex-noise · d3-delaunay"),
    ("Grammar service", "LanguageTool"),
    ("Storage", "Local filesystem — plain JSON + assets, offline-first (no cloud, no account)"),
    ("Document formats", "Notes (HTML) · Mind maps (JSON) · Fantasy maps (JSON)"),
    ("Export formats", "PDF (books) · PNG (maps)"),
    ("Import formats", "PNG · SVG · WebP · JPG (map assets & base images)"),
    ("License", "Proprietary — all rights reserved"),
]
st = doc.add_table(rows=0, cols=2)
st.alignment = WD_TABLE_ALIGNMENT.CENTER
set_col_widths(st, [Inches(1.95), Inches(4.95)])
for i, (k, v) in enumerate(specs):
    cells = st.add_row().cells
    for c in cells:
        cell_margins(c)
    shade(cells[0], LIGHT_HEX)
    if i % 2 == 1:
        shade(cells[1], "F6F7FE")
    run(cells[0].paragraphs[0], k, size=9.5, bold=True, color=INK)
    run(cells[1].paragraphs[0], v, size=9.5, color=INK)
table_borders(st, inside=True)

# ── footer note ────────────────────────────────────────────────────────────
doc.add_paragraph().paragraph_format.space_after = Pt(2)
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(fp, "MnemoScript © 2026 · All rights reserved.   "
        "The hand-drawn ink icon set is original work; remaining map glyphs are from "
        "game-icons.net (CC BY 3.0).", size=8, italic=True, color=MUTED)

import os
out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "MnemoScript_Datasheet.docx")
doc.save(out)
print("Saved:", out)
